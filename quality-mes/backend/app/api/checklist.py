from fastapi import APIRouter, Depends, HTTPException, Query, status, UploadFile, File
from sqlalchemy.orm import Session

from app.auth.security import get_current_user
from app.core.database import get_db
from app.models.checklist import ChecklistTemplate

router = APIRouter(prefix="/api/checklists", tags=["Checklists"])


@router.get("/")
def list_templates(
    module: str | None = Query(None),
    db: Session = Depends(get_db),
    _token: dict = Depends(get_current_user),
):
    query = db.query(ChecklistTemplate)
    if module:
        query = query.filter(ChecklistTemplate.module == module)
    templates = query.order_by(ChecklistTemplate.name).all()
    return [
        {
            "id": t.id, "code": t.code, "name": t.name, "module": t.module,
            "items": t.items, "created_at": str(t.created_at),
        }
        for t in templates
    ]


@router.post("/", status_code=status.HTTP_201_CREATED)
def create_template(data: dict, db: Session = Depends(get_db), _token: dict = Depends(get_current_user)):
    existing = db.query(ChecklistTemplate).filter(ChecklistTemplate.code == data["code"]).first()
    if existing:
        raise HTTPException(status_code=400, detail="Ma checklist da ton tai")
    if data["module"] not in ("iqc", "oqc", "ipqc"):
        raise HTTPException(status_code=400, detail="Module khong hop le")
    items = data.get("items", [])
    if isinstance(items, str):
        import json
        items = json.loads(items)
    t = ChecklistTemplate(code=data["code"], name=data["name"], module=data["module"], items=items)
    db.add(t)
    db.commit()
    db.refresh(t)
    return {"id": t.id, "code": t.code, "name": t.name, "module": t.module, "items": t.items}


@router.put("/{template_id}")
def update_template(template_id: int, data: dict, db: Session = Depends(get_db), _token: dict = Depends(get_current_user)):
    t = db.query(ChecklistTemplate).filter(ChecklistTemplate.id == template_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Khong tim thay checklist")
    if "name" in data:
        t.name = data["name"]
    items = data.get("items")
    if items is not None:
        if isinstance(items, str):
            import json
            items = json.loads(items)
        t.items = items
    db.commit()
    return {"message": "Cap nhat thanh cong"}


@router.delete("/{template_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_template(template_id: int, db: Session = Depends(get_db), _token: dict = Depends(get_current_user)):
    t = db.query(ChecklistTemplate).filter(ChecklistTemplate.id == template_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Khong tim thay checklist")
    db.delete(t)
    db.commit()


@router.post("/import-excel")
def import_checklist_excel(
    module: str = Query(..., pattern="^(iqc|oqc|ipqc)$"),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    _token: dict = Depends(get_current_user),
):
    from io import BytesIO

    contents = file.file.read()
    filename = file.filename or ""

    if filename.lower().endswith(".docx"):
        return _import_from_docx(contents, module, db)
    else:
        return _import_from_xlsx(contents, module, db)


def _import_from_xlsx(contents: bytes, module: str, db: Session):
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(contents))
    count = 0

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        code = str(sheet_name).strip()
        code = code.replace(" ", "-").replace("/", "-")
        name = str(sheet_name).strip()
        items = []

        for row in ws.iter_rows(min_row=1, values_only=True):
            if not row or not row[0]:
                continue
            item_name = str(row[0]).strip()
            if item_name.lower() in ("muc kiem", "item", "ten muc", "no", "mục kiểm", "tiêu chuẩn"):
                continue
            spec = str(row[1]).strip() if len(row) > 1 and row[1] else ""
            try:
                min_v = float(row[2]) if len(row) > 2 and row[2] is not None else None
            except: min_v = None
            try:
                max_v = float(row[3]) if len(row) > 3 and row[3] is not None else None
            except: max_v = None

            items.append({
                "item_name": item_name,
                "specification": spec,
                "standard_min": min_v,
                "standard_max": max_v,
            })

        if items:
            existing = db.query(ChecklistTemplate).filter(ChecklistTemplate.code == code).first()
            if not existing:
                db.add(ChecklistTemplate(code=code, name=name, module=module, items=items))
                count += 1

    db.commit()
    return {"message": f"Da import {count} checklist tu file Excel"}


def _import_from_docx(contents: bytes, module: str, db: Session):
    from docx import Document

    doc = Document(BytesIO(contents))
    count = 0

    for idx, table in enumerate(doc.tables):
        rows = table.rows
        if len(rows) < 2:
            continue

        text_before = ""
        for para in doc.paragraphs:
            ptext = para.text.strip()
            if ptext and ptext not in text_before:
                if not text_before:
                    text_before = ptext
                break

        code = f"{module.upper()}-CL-{idx+1:03d}"
        name = text_before if text_before else f"Checklist {idx+1}"
        items = []

        for row in rows:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells or not cells[0]:
                continue
            item_name = cells[0]
            if item_name.lower() in ("muc kiem", "item", "ten muc", "no", "mục kiểm", "tiêu chuẩn", "stt"):
                continue
            spec = cells[1] if len(cells) > 1 else ""
            try:
                min_v = float(cells[2]) if len(cells) > 2 and cells[2] else None
            except: min_v = None
            try:
                max_v = float(cells[3]) if len(cells) > 3 and cells[3] else None
            except: max_v = None

            items.append({
                "item_name": item_name,
                "specification": spec,
                "standard_min": min_v,
                "standard_max": max_v,
            })

        if items:
            name_slug = name[:30].replace(" ", "-").replace("/", "-")
            code = f"{module.upper()}-{name_slug}"
            existing = db.query(ChecklistTemplate).filter(ChecklistTemplate.code == code).first()
            if not existing:
                db.add(ChecklistTemplate(code=code, name=name, module=module, items=items))
                count += 1

    db.commit()
    return {"message": f"Da import {count} checklist tu file Word"}
